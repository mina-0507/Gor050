# gui_app.py - с графиком для рассрочки

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from bank_package import CreditCalculator, InstallmentCalculator, DepositCalculator
from openpyxl import Workbook
from docx import Document

class BankApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Банковские услуги")
        self.root.geometry("600x650")
        
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True, padx=5, pady=5)
        
        self.credit_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.credit_tab, text="Кредит")
        self.setup_credit()
        
        self.inst_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.inst_tab, text="Рассрочка")
        self.setup_installment()
        
        self.deposit_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.deposit_tab, text="Вклад")
        self.setup_deposit()
        
        self.root.mainloop()
    
    def setup_credit(self):
        frame = ttk.LabelFrame(self.credit_tab, text="Введите данные", padding=10)
        frame.pack(fill='x', padx=10, pady=10)
        
        ttk.Label(frame, text="Сумма кредита (руб):").grid(row=0, column=0, pady=5)
        self.credit_sum = tk.Entry(frame, width=20)
        self.credit_sum.insert(0, "100000")
        self.credit_sum.grid(row=0, column=1, pady=5)
        
        ttk.Label(frame, text="Ставка (%):").grid(row=1, column=0, pady=5)
        self.credit_rate = tk.Entry(frame, width=20)
        self.credit_rate.insert(0, "15")
        self.credit_rate.grid(row=1, column=1, pady=5)
        
        ttk.Label(frame, text="Срок (мес):").grid(row=2, column=0, pady=5)
        self.credit_months = tk.Entry(frame, width=20)
        self.credit_months.insert(0, "12")
        self.credit_months.grid(row=2, column=1, pady=5)
        
        btn = tk.Button(frame, text="Рассчитать", bg='red', fg='white', command=self.calc_credit)
        btn.grid(row=3, column=0, columnspan=2, pady=10)
        
        res_frame = ttk.LabelFrame(self.credit_tab, text="Результат", padding=10)
        res_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        text_frame = tk.Frame(res_frame)
        text_frame.pack(fill='both', expand=True)
        
        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.pack(side='right', fill='y')
        
        self.credit_text = tk.Text(text_frame, height=12, width=60, yscrollcommand=scrollbar.set)
        self.credit_text.pack(side='left', fill='both', expand=True)
        scrollbar.config(command=self.credit_text.yview)
        
        btn_frame = tk.Frame(res_frame)
        btn_frame.pack(fill='x', pady=5)
        
        btn_excel = tk.Button(btn_frame, text="Сохранить в Excel", bg='green', fg='white', 
                              command=lambda: self.save_excel('credit'))
        btn_excel.pack(side='left', padx=5)
        
        btn_word = tk.Button(btn_frame, text="Сохранить в Word", bg='blue', fg='white',
                             command=lambda: self.save_word('credit'))
        btn_word.pack(side='left', padx=5)
    
    def calc_credit(self):
        try:
            s = float(self.credit_sum.get())
            r = float(self.credit_rate.get())
            m = int(self.credit_months.get())
            
            calc = CreditCalculator(s, r, m)
            schedule = calc.get_schedule_annuity()
            
            self.credit_text.delete(1.0, tk.END)
            
            self.credit_text.insert(tk.END, "="*50 + "\n")
            self.credit_text.insert(tk.END, "ОСНОВНЫЕ ИТОГИ\n")
            self.credit_text.insert(tk.END, "="*50 + "\n")
            self.credit_text.insert(tk.END, f"Сумма кредита:   {s:>10,.2f} руб\n")
            self.credit_text.insert(tk.END, f"Ставка:          {r:>10} %\n")
            self.credit_text.insert(tk.END, f"Срок:            {m:>10} мес\n")
            self.credit_text.insert(tk.END, f"Ежемес. платёж:  {calc.monthly_payment():>10,.2f} руб\n")
            self.credit_text.insert(tk.END, f"Общая сумма:     {calc.total_payment():>10,.2f} руб\n")
            self.credit_text.insert(tk.END, f"Переплата:       {calc.overpayment():>10,.2f} руб\n")
            
            self.credit_text.insert(tk.END, "\n" + "="*50 + "\n")
            self.credit_text.insert(tk.END, "ГРАФИК ПЛАТЕЖЕЙ\n")
            self.credit_text.insert(tk.END, "="*50 + "\n")
            self.credit_text.insert(tk.END, f"{'Месяц':<6} {'Платёж':<12} {'Проценты':<10} {'Остаток':<12}\n")
            self.credit_text.insert(tk.END, "-"*50 + "\n")
            
            for p in schedule:
                self.credit_text.insert(tk.END, f"{p['month']:<6} {p['payment']:>10,.2f} {p['interest']:>9,.2f} {p['remaining']:>11,.2f}\n")
            
            self.credit_data = {
                'sum': s, 'rate': r, 'months': m,
                'payment': calc.monthly_payment(),
                'total': calc.total_payment(),
                'over': calc.overpayment(),
                'schedule': schedule
            }
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка: {e}")
    
    def setup_installment(self):
        frame = ttk.LabelFrame(self.inst_tab, text="Введите данные", padding=10)
        frame.pack(fill='x', padx=10, pady=10)
        
        ttk.Label(frame, text="Сумма покупки (руб):").grid(row=0, column=0, pady=5)
        self.inst_sum = tk.Entry(frame, width=20)
        self.inst_sum.insert(0, "50000")
        self.inst_sum.grid(row=0, column=1, pady=5)
        
        ttk.Label(frame, text="Срок (мес):").grid(row=1, column=0, pady=5)
        self.inst_months = tk.Entry(frame, width=20)
        self.inst_months.insert(0, "6")
        self.inst_months.grid(row=1, column=1, pady=5)
        
        btn = tk.Button(frame, text="Рассчитать", bg='red', fg='white', command=self.calc_installment)
        btn.grid(row=2, column=0, columnspan=2, pady=10)
        
        res_frame = ttk.LabelFrame(self.inst_tab, text="Результат", padding=10)
        res_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        text_frame = tk.Frame(res_frame)
        text_frame.pack(fill='both', expand=True)
        
        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.pack(side='right', fill='y')
        
        self.inst_text = tk.Text(text_frame, height=12, width=60, yscrollcommand=scrollbar.set)
        self.inst_text.pack(side='left', fill='both', expand=True)
        scrollbar.config(command=self.inst_text.yview)
        
        btn_frame = tk.Frame(res_frame)
        btn_frame.pack(fill='x', pady=5)
        
        btn_excel = tk.Button(btn_frame, text="Сохранить в Excel", bg='green', fg='white',
                              command=lambda: self.save_excel('installment'))
        btn_excel.pack(side='left', padx=5)
        
        btn_word = tk.Button(btn_frame, text="Сохранить в Word", bg='blue', fg='white',
                             command=lambda: self.save_word('installment'))
        btn_word.pack(side='left', padx=5)
    
    def calc_installment(self):
        try:
            s = float(self.inst_sum.get())
            m = int(self.inst_months.get())
            
            calc = InstallmentCalculator(s, m)
            payment = calc.monthly_payment()
            
            # Строим график платежей
            schedule = []
            remaining = s
            for i in range(1, m + 1):
                remaining = remaining - payment
                schedule.append({
                    'month': i,
                    'payment': payment,
                    'remaining': round(max(remaining, 0), 2)
                })
            
            self.inst_text.delete(1.0, tk.END)
            
            self.inst_text.insert(tk.END, "="*50 + "\n")
            self.inst_text.insert(tk.END, "ОСНОВНЫЕ ИТОГИ\n")
            self.inst_text.insert(tk.END, "="*50 + "\n")
            self.inst_text.insert(tk.END, f"Сумма покупки:   {s:>10,.2f} руб\n")
            self.inst_text.insert(tk.END, f"Срок:            {m:>10} мес\n")
            self.inst_text.insert(tk.END, f"Ежемес. платёж:  {payment:>10,.2f} руб\n")
            self.inst_text.insert(tk.END, f"Общая сумма:     {s:>10,.2f} руб\n")
            self.inst_text.insert(tk.END, f"Переплата:             0.00 руб\n")
            
            self.inst_text.insert(tk.END, "\n" + "="*50 + "\n")
            self.inst_text.insert(tk.END, "ГРАФИК ПЛАТЕЖЕЙ\n")
            self.inst_text.insert(tk.END, "="*50 + "\n")
            self.inst_text.insert(tk.END, f"{'Месяц':<6} {'Платёж':<12} {'Остаток':<12}\n")
            self.inst_text.insert(tk.END, "-"*50 + "\n")
            
            for p in schedule:
                self.inst_text.insert(tk.END, f"{p['month']:<6} {p['payment']:>10,.2f} {p['remaining']:>11,.2f}\n")
            
            self.inst_data = {
                'sum': s, 'months': m, 'payment': payment,
                'schedule': schedule
            }
        except:
            messagebox.showerror("Ошибка", "Введи числа правильно!")
    
    def setup_deposit(self):
        frame = ttk.LabelFrame(self.deposit_tab, text="Введите данные", padding=10)
        frame.pack(fill='x', padx=10, pady=10)
        
        ttk.Label(frame, text="Сумма вклада (руб):").grid(row=0, column=0, pady=5)
        self.dep_sum = tk.Entry(frame, width=20)
        self.dep_sum.insert(0, "100000")
        self.dep_sum.grid(row=0, column=1, pady=5)
        
        ttk.Label(frame, text="Ставка (%):").grid(row=1, column=0, pady=5)
        self.dep_rate = tk.Entry(frame, width=20)
        self.dep_rate.insert(0, "12")
        self.dep_rate.grid(row=1, column=1, pady=5)
        
        ttk.Label(frame, text="Срок (мес):").grid(row=2, column=0, pady=5)
        self.dep_months = tk.Entry(frame, width=20)
        self.dep_months.insert(0, "12")
        self.dep_months.grid(row=2, column=1, pady=5)
        
        btn = tk.Button(frame, text="Рассчитать", bg='red', fg='white', command=self.calc_deposit)
        btn.grid(row=3, column=0, columnspan=2, pady=10)
        
        res_frame = ttk.LabelFrame(self.deposit_tab, text="Результат", padding=10)
        res_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        text_frame = tk.Frame(res_frame)
        text_frame.pack(fill='both', expand=True)
        
        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.pack(side='right', fill='y')
        
        self.dep_text = tk.Text(text_frame, height=12, width=60, yscrollcommand=scrollbar.set)
        self.dep_text.pack(side='left', fill='both', expand=True)
        scrollbar.config(command=self.dep_text.yview)
        
        btn_frame = tk.Frame(res_frame)
        btn_frame.pack(fill='x', pady=5)
        
        btn_excel = tk.Button(btn_frame, text="Сохранить в Excel", bg='green', fg='white',
                              command=lambda: self.save_excel('deposit'))
        btn_excel.pack(side='left', padx=5)
        
        btn_word = tk.Button(btn_frame, text="Сохранить в Word", bg='blue', fg='white',
                             command=lambda: self.save_word('deposit'))
        btn_word.pack(side='left', padx=5)
    
    def calc_deposit(self):
        try:
            s = float(self.dep_sum.get())
            r = float(self.dep_rate.get())
            m = int(self.dep_months.get())
            
            calc = DepositCalculator(s, r, m)
            schedule = calc.get_schedule()
            
            self.dep_text.delete(1.0, tk.END)
            
            self.dep_text.insert(tk.END, "="*50 + "\n")
            self.dep_text.insert(tk.END, "ОСНОВНЫЕ ИТОГИ\n")
            self.dep_text.insert(tk.END, "="*50 + "\n")
            self.dep_text.insert(tk.END, f"Сумма вклада:    {s:>10,.2f} руб\n")
            self.dep_text.insert(tk.END, f"Ставка:          {r:>10} %\n")
            self.dep_text.insert(tk.END, f"Срок:            {m:>10} мес\n")
            self.dep_text.insert(tk.END, f"Итоговая сумма:  {calc.final_amount():>10,.2f} руб\n")
            self.dep_text.insert(tk.END, f"Начислено %:     {calc.total_interest():>10,.2f} руб\n")
            
            self.dep_text.insert(tk.END, "\n" + "="*50 + "\n")
            self.dep_text.insert(tk.END, "ГРАФИК НАЧИСЛЕНИЯ\n")
            self.dep_text.insert(tk.END, "="*50 + "\n")
            self.dep_text.insert(tk.END, f"{'Месяц':<6} {'Сумма на счете':<15} {'Начислено %':<12}\n")
            self.dep_text.insert(tk.END, "-"*50 + "\n")
            
            for p in schedule:
                self.dep_text.insert(tk.END, f"{p['month']:<6} {p['amount']:>13,.2f} {p['interest']:>11,.2f}\n")
            
            self.dep_data = {
                'sum': s, 'rate': r, 'months': m,
                'final': calc.final_amount(),
                'interest': calc.total_interest(),
                'schedule': schedule
            }
        except:
            messagebox.showerror("Ошибка", "Введи числа правильно!")
    
    def save_excel(self, typ):
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Банковский расчёт"
            
            if typ == 'credit':
                d = self.credit_data
                ws['A1'] = "КРЕДИТ"
                ws['A2'] = f"Сумма: {d['sum']} руб"
                ws['A3'] = f"Ставка: {d['rate']}%"
                ws['A4'] = f"Срок: {d['months']} мес"
                ws['A5'] = f"Платёж: {d['payment']} руб"
                ws['A6'] = f"Всего: {d['total']} руб"
                ws['A7'] = f"Переплата: {d['over']} руб"
                
                ws['A9'] = "Месяц"
                ws['B9'] = "Платёж"
                ws['C9'] = "Проценты"
                ws['D9'] = "Остаток"
                row = 10
                for p in d['schedule']:
                    ws[f'A{row}'] = p['month']
                    ws[f'B{row}'] = p['payment']
                    ws[f'C{row}'] = p['interest']
                    ws[f'D{row}'] = p['remaining']
                    row += 1
                    
            elif typ == 'installment':
                d = self.inst_data
                ws['A1'] = "РАССРОЧКА"
                ws['A2'] = f"Сумма: {d['sum']} руб"
                ws['A3'] = f"Срок: {d['months']} мес"
                ws['A4'] = f"Платёж: {d['payment']} руб"
                ws['A5'] = "Переплата: 0 руб"
                
                # График платежей для рассрочки
                ws['A7'] = "Месяц"
                ws['B7'] = "Платёж"
                ws['C7'] = "Остаток"
                row = 8
                for p in d['schedule']:
                    ws[f'A{row}'] = p['month']
                    ws[f'B{row}'] = p['payment']
                    ws[f'C{row}'] = p['remaining']
                    row += 1
            
            else:
                d = self.dep_data
                ws['A1'] = "ВКЛАД"
                ws['A2'] = f"Сумма: {d['sum']} руб"
                ws['A3'] = f"Ставка: {d['rate']}%"
                ws['A4'] = f"Срок: {d['months']} мес"
                ws['A5'] = f"Итого: {d['final']} руб"
                ws['A6'] = f"Проценты: {d['interest']} руб"
                
                ws['A8'] = "Месяц"
                ws['B8'] = "Сумма на счете"
                ws['C8'] = "Начислено %"
                row = 9
                for p in d['schedule']:
                    ws[f'A{row}'] = p['month']
                    ws[f'B{row}'] = p['amount']
                    ws[f'C{row}'] = p['interest']
                    row += 1
            
            path = filedialog.asksaveasfilename(defaultextension=".xlsx")
            if path:
                wb.save(path)
                messagebox.showinfo("Успех", "Файл сохранён!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не сохранилось: {e}")
    
    def save_word(self, typ):
        try:
            doc = Document()
            doc.add_heading('Банковский расчёт', 0)
            
            if typ == 'credit':
                d = self.credit_data
                doc.add_heading('Кредит', level=1)
                doc.add_paragraph(f'Сумма: {d["sum"]:,.2f} руб')
                doc.add_paragraph(f'Ставка: {d["rate"]}%')
                doc.add_paragraph(f'Срок: {d["months"]} мес')
                doc.add_paragraph(f'Платёж: {d["payment"]:,.2f} руб')
                doc.add_paragraph(f'Всего: {d["total"]:,.2f} руб')
                doc.add_paragraph(f'Переплата: {d["over"]:,.2f} руб')
                
                doc.add_heading('График платежей', level=2)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                headers = ['Месяц', 'Платёж', 'Проценты', 'Остаток']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for p in d['schedule']:
                    row = table.add_row()
                    row.cells[0].text = str(p['month'])
                    row.cells[1].text = f"{p['payment']:,.2f}"
                    row.cells[2].text = f"{p['interest']:,.2f}"
                    row.cells[3].text = f"{p['remaining']:,.2f}"
                    
            elif typ == 'installment':
                d = self.inst_data
                doc.add_heading('Рассрочка', level=1)
                doc.add_paragraph(f'Сумма: {d["sum"]:,.2f} руб')
                doc.add_paragraph(f'Срок: {d["months"]} мес')
                doc.add_paragraph(f'Платёж: {d["payment"]:,.2f} руб')
                doc.add_paragraph('Переплата: 0 руб')
                
                # График для рассрочки
                doc.add_heading('График платежей', level=2)
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                headers = ['Месяц', 'Платёж', 'Остаток']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for p in d['schedule']:
                    row = table.add_row()
                    row.cells[0].text = str(p['month'])
                    row.cells[1].text = f"{p['payment']:,.2f}"
                    row.cells[2].text = f"{p['remaining']:,.2f}"
            
            else:
                d = self.dep_data
                doc.add_heading('Вклад', level=1)
                doc.add_paragraph(f'Сумма: {d["sum"]:,.2f} руб')
                doc.add_paragraph(f'Ставка: {d["rate"]}%')
                doc.add_paragraph(f'Срок: {d["months"]} мес')
                doc.add_paragraph(f'Итого: {d["final"]:,.2f} руб')
                doc.add_paragraph(f'Проценты: {d["interest"]:,.2f} руб')
                
                doc.add_heading('График начисления', level=2)
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                headers = ['Месяц', 'Сумма на счете', 'Начислено %']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for p in d['schedule']:
                    row = table.add_row()
                    row.cells[0].text = str(p['month'])
                    row.cells[1].text = f"{p['amount']:,.2f}"
                    row.cells[2].text = f"{p['interest']:,.2f}"
            
            path = filedialog.asksaveasfilename(defaultextension=".docx")
            if path:
                doc.save(path)
                messagebox.showinfo("Успех", "Файл сохранён!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не сохранилось: {e}")

if __name__ == "__main__":
    BankApp()